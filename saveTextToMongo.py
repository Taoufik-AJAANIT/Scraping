# -*- coding: utf-8 -*-
import os
import pytesseract
from PIL import Image
from docx import Document
from wand.image import Image as wi
from mongo import Mongo





def toText(path='path to Offres Directorie'):
    text = ''

    #set up database
    mydb = Mongo('TTProject', ['offres-info', 'offres-text'])

    #move to directory that containes offres
    os.chdir(path)
    dirnames = open('names.txt').read().split(" ")
    references = open('references.txt').read().split(" ")
    i = 0
    for directory in dirnames:
        if len(directory):
            x = len(directory)
            dirname = str(directory)[1:x-1]
    
            #informations to store
            reference = references[i]
            name = dirname
            text += '############################  ' + \
                dirname+'   ############################\n\n\n'
            #loop trought Directory :
            dirPath = path+'\\'+dirname
            for filename in os.listdir(os.path.join(dirPath)):
                text += '\n\n####################   '+filename+'   ####################\n\n'
                filePath = path+'\\'+dirname+'\\'+filename
                # check extention
                # PDFs
                if filename.endswith('.pdf'):
                    pdfFile = wi(filename=os.path.join(filePath), resolution=300)
                    images = pdfFile.convert('jpeg')
                    for page in images:
                        page.save('page.jpg', 'JPEG')
                        image = Image.open(io.BytesIO(page))
                        imgPage = wi(image=image)
                        Image = imgPage.make_blob('jpeg')
                        text += pytesseract.image_to_string(
                            Image.open('page.jpg'), lang='fra')
                # DOCXs
                if filename.endswith('.docx'):
                    document = Document(filePath)
                    tables = document.tables
                    for p in document.paragraphs:
                        text += p.text
                    text += "\n  ###### Tables's Content ###### \n"
                    for table in tables:
                        for row in table.rows:
                            for cell in row.cells:
                                for paragraph in cell.paragraphs:
                                      text += (paragraph.text)
                                      
                #inserting to mongo
                offre = {'_id' : reference , "name" : name , 'text' : text}
                mydb.insert(offre , 'offres-text')
                text = ''
                i += 1
